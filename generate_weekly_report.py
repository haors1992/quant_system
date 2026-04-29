"""生成周报 Word 文档
使用方式：
  1. 安装依赖：pip install python-docx
  2. 运行脚本：python generate_weekly_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)


def add_table_row(table, cells_data):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return row


def create_weekly_report(output_path):
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ===== 标题 =====
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('周  报')
    run.font.size = Pt(22)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.bold = True

    # 副标题 - 时间
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（2026年4月27日 — 4月30日）')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 基本信息表
    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('姓名：', '__________', '岗位：', '研发工程师'),
        ('部门：', '__________', '日期：', '2026年4月30日'),
    ]

    for row_idx, row_data in enumerate(info_data):
        for col_idx, text in enumerate(row_data):
            cell = info_table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if col_idx % 2 == 0:
                run.bold = True
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()

    # ===== 一、本周工作完成情况 =====
    h1 = doc.add_heading('', level=1)
    run = h1.add_run('一、本周工作完成情况')
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # 表格
    work_table = doc.add_table(rows=1, cols=3)
    work_table.style = 'Table Grid'
    work_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['序号', '工作内容', '状态']
    for i, text in enumerate(headers):
        cell = work_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = True
        set_cell_shading(cell, 'D9E2F3')

    for row in work_table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(11)
        row.cells[2].width = Cm(2)

    work_items = [
        ('1', '4月需求开发', '进行中'),
        ('2', '编排中心-集客资源准备单工程编码关联资源查询功能优化', '已完成'),
        ('3', '【集客】集客专线资源选择增加光缆施工接入点设计信息继承回填', '已完成'),
        ('4', '宽带电视同开电视单出库取消ONU处理需求', '已完成'),
        ('5', '【APP】家宽开通同装电视单取消ONU信息校验需求', '已完成'),
        ('6', '智能体开发', '进行中'),
    ]

    for seq, content, status in work_items:
        add_table_row(work_table, [seq, content, status])

    doc.add_paragraph()

    # 详细说明
    h2 = doc.add_heading('', level=2)
    run = h2.add_run('详细说明')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    details = [
        ('1. 4月需求开发', '参与4月需求开发工作，按计划推进各项任务。'),
        ('2. 编排中心-集客资源准备单工程编码关联资源查询功能优化', '完成编排中心集客资源准备单模块中工程编码关联资源查询功能的优化，提升查询效率和准确性。'),
        ('3. 【集客】集客专线资源选择增加光缆施工接入点设计信息继承回填', '在集客专线资源选择流程中增加光缆施工接入点的设计信息继承与回填功能，完善资源数据流转。'),
        ('4. 宽带电视同开电视单出库取消ONU处理需求', '完成宽带与电视同开场景下，电视单出库时取消ONU处理的流程改造。'),
        ('5. 【APP】家宽开通同装电视单取消ONU信息校验需求', '实现家宽开通同装电视场景下，电视单取消时ONU信息的校验逻辑，确保数据一致性。'),
        ('6. 智能体开发', '参与智能体开发工作，持续推进功能设计与实现。'),
    ]

    for title_text, desc in details:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        run = p2.add_run(desc)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # ===== 二、下周工作计划 =====
    h1 = doc.add_heading('', level=1)
    run = h1.add_run('二、下周工作计划')
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.style = 'Table Grid'
    plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    plan_headers = ['序号', '计划内容', '预计完成时间']
    for i, text in enumerate(plan_headers):
        cell = plan_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = True
        set_cell_shading(cell, 'D9E2F3')

    for row in plan_table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(10)
        row.cells[2].width = Cm(3)

    plan_items = [
        ('1', '4月需求开发收尾', '4月30日'),
        ('2', '智能体开发持续迭代', '进行中'),
        ('3', '待补充', '-'),
    ]

    for seq, content, deadline in plan_items:
        add_table_row(plan_table, [seq, content, deadline])

    doc.add_paragraph()

    # ===== 三、问题与风险 =====
    h1 = doc.add_heading('', level=1)
    run = h1.add_run('三、问题与风险')
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p = doc.add_paragraph()
    run = p.add_run('无')
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # ===== 四、需要协调的事项 =====
    h1 = doc.add_heading('', level=1)
    run = h1.add_run('四、需要协调的事项')
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p = doc.add_paragraph()
    run = p.add_run('无')
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ===== 保存 =====
    doc.save(output_path)
    print(f'周报已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output = os.path.join(script_dir, '周报_2026年4月第4周.docx')
    create_weekly_report(output)
